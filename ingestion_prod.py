"""Production-ready document ingestion with comprehensive error handling"""
import os
import re
import time
from pathlib import Path
from typing import List, Dict, Optional, Any

import PyPDF2
from docx import Document

from config import config
from production_logger import get_logger, log_file_upload
from validation import validate_file_content, ValidationError
from exception_handler import handle_exceptions, retry_on_failure, FileProcessingError, file_circuit_breaker

logger = get_logger('ingestion')

class DocumentProcessor:
    """Production-ready document processing class"""
    
    def __init__(self):
        self.supported_types = {
            '.pdf': self._extract_text_from_pdf,
            '.docx': self._extract_text_from_docx,
            '.txt': self._extract_text_from_txt
        }
        self.processed_files: Dict[str, Dict[str, Any]] = {}
    
    @handle_exceptions("extract_text_from_pdf", default_return="")
    @retry_on_failure(max_retries=2, delay=1.0)
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Enhanced PDF text extraction with better error handling"""
        logger.info(f"Processing PDF: {file_path.name}")
        
        try:
            text_content = []
            
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                
                # Validate PDF structure
                if len(reader.pages) == 0:
                    raise FileProcessingError("PDF contains no pages", str(file_path))
                
                if len(reader.pages) > 1000:  # Security: prevent processing huge files
                    logger.warning(f"PDF has {len(reader.pages)} pages, limiting to first 1000")
                    pages_to_process = reader.pages[:1000]
                else:
                    pages_to_process = reader.pages
                
                for page_num, page in enumerate(pages_to_process):
                    try:
                        page_text = page.extract_text()
                        
                        if page_text and page_text.strip():
                            # Clean up common PDF extraction issues
                            cleaned_text = self._clean_pdf_text(page_text)
                            
                            if cleaned_text and len(cleaned_text.strip()) > 10:
                                text_content.append(f"\n--- Page {page_num + 1} ---\n{cleaned_text}")
                                
                        # Security: prevent memory issues with very large documents
                        if len(''.join(text_content)) > 10 * 1024 * 1024:  # 10MB text limit
                            logger.warning(f"PDF text extraction stopped at page {page_num + 1} due to size limit")
                            break
                            
                    except Exception as page_error:
                        logger.warning(f"Error extracting page {page_num + 1} from {file_path.name}: {page_error}")
                        continue
            
            final_text = '\n'.join(text_content).strip()
            
            if not final_text:
                raise FileProcessingError("No readable text found in PDF", str(file_path))
            
            logger.info(f"Successfully extracted {len(final_text)} characters from {file_path.name}")
            return final_text
            
        except PyPDF2.errors.PdfReadError as e:
            raise FileProcessingError(f"Invalid or corrupted PDF: {e}", str(file_path))
        except Exception as e:
            raise FileProcessingError(f"PDF processing failed: {e}", str(file_path))
    
    def _clean_pdf_text(self, text: str) -> str:
        """Clean PDF text extraction artifacts"""
        if not text:
            return ""
        
        # Remove null bytes and control characters (security)
        text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\t\r')
        
        # Fix common PDF extraction issues
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)  # Add space between camelCase
        text = re.sub(r'(\w)(\d)', r'\1 \2', text)  # Add space between word and number
        text = re.sub(r'(\d)([A-Za-z])', r'\1 \2', text)  # Add space between number and word
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)  # Remove excessive line breaks
        
        return text.strip()
    
    @handle_exceptions("extract_text_from_docx", default_return="")
    @retry_on_failure(max_retries=2, delay=1.0)
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Enhanced DOCX text extraction"""
        logger.info(f"Processing DOCX: {file_path.name}")
        
        try:
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraph text
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            final_text = '\n'.join(text_parts)
            
            if not final_text.strip():
                raise FileProcessingError("No readable text found in DOCX", str(file_path))
            
            # Security: limit text size
            if len(final_text) > 10 * 1024 * 1024:  # 10MB limit
                final_text = final_text[:10 * 1024 * 1024]
                logger.warning(f"DOCX text truncated to 10MB for {file_path.name}")
            
            logger.info(f"Successfully extracted {len(final_text)} characters from {file_path.name}")
            return final_text
            
        except Exception as e:
            raise FileProcessingError(f"DOCX processing failed: {e}", str(file_path))
    
    @handle_exceptions("extract_text_from_txt", default_return="")
    def _extract_text_from_txt(self, file_path: Path) -> str:
        """Enhanced TXT file processing with encoding detection"""
        logger.info(f"Processing TXT: {file_path.name}")
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        
                        # Validate content
                        if content.strip():
                            # Security: limit file size
                            if len(content) > 10 * 1024 * 1024:  # 10MB limit
                                content = content[:10 * 1024 * 1024]
                                logger.warning(f"TXT file truncated to 10MB for {file_path.name}")
                            
                            logger.info(f"Successfully read {len(content)} characters from {file_path.name} using {encoding}")
                            return content
                            
                except UnicodeDecodeError:
                    continue
            
            raise FileProcessingError("Could not decode text file with any supported encoding", str(file_path))
            
        except Exception as e:
            raise FileProcessingError(f"TXT processing failed: {e}", str(file_path))
    
    @handle_exceptions("chunk_text", default_return=[])
    def chunk_text(self, text: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """Split text into overlapping chunks with enhanced logic"""
        chunk_size = chunk_size or config.CHUNK_SIZE
        overlap = overlap or config.CHUNK_OVERLAP
        
        if not text or len(text) <= chunk_size:
            return [text] if text.strip() else []
        
        chunks = []
        start = 0
        text_length = len(text)
        
        while start < text_length:
            end = min(start + chunk_size, text_length)
            
            # Try to break at sentence boundary
            if end < text_length:
                # Look for sentence endings near the chunk boundary
                search_start = max(start + chunk_size - 200, start + 1)
                for i in range(end, search_start, -1):
                    if i > 0 and text[i-1:i+1] in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                        end = i
                        break
                else:
                    # If no sentence boundary found, try paragraph breaks
                    for i in range(end, search_start, -1):
                        if i > 0 and text[i-1:i+1] in ['\n\n', '\r\n\r\n']:
                            end = i
                            break
            
            chunk = text[start:end].strip()
            
            # Only add substantial chunks
            if chunk and len(chunk) > config.MIN_CHUNK_SIZE:
                chunks.append(chunk)
            
            # Calculate next start position
            if end >= text_length:
                break
            
            start = max(end - overlap, start + 1)
        
        logger.info(f"Created {len(chunks)} chunks from {len(text)} characters")
        return chunks
    
    @handle_exceptions("process_file", default_return=None)
    def process_file(self, file_path: Path) -> Optional[Dict[str, Any]]:
        """Process a single file with comprehensive validation"""
        start_time = time.time()
        
        try:
            # Validate file
            file_info = validate_file_content(file_path)
            
            # Check file extension
            file_ext = file_path.suffix.lower()
            if file_ext not in self.supported_types:
                raise FileProcessingError(f"Unsupported file type: {file_ext}", str(file_path))
            
            # Use circuit breaker for file processing
            def extract_text():
                return self.supported_types[file_ext](file_path)
            
            text_content = file_circuit_breaker.call(extract_text)
            
            if not text_content or not text_content.strip():
                raise FileProcessingError("No readable content extracted", str(file_path))
            
            # Validate extracted text
            if len(text_content) < 50:
                logger.warning(f"Very short content extracted from {file_path.name}: {len(text_content)} characters")
            
            processing_time = time.time() - start_time
            
            # Log successful processing
            log_file_upload(
                filename=file_path.name,
                file_size=file_info['size'],
                processing_time=processing_time
            )
            
            result = {
                'filename': file_path.name,
                'content': text_content,
                'original_filename': file_path.name,
                'file_info': file_info,
                'processing_time': processing_time,
                'content_length': len(text_content)
            }
            
            self.processed_files[file_path.name] = result
            logger.info(f"Successfully processed {file_path.name} in {processing_time:.2f}s")
            
            return result
            
        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"Failed to process {file_path.name} after {processing_time:.2f}s: {e}")
            raise

@handle_exceptions("load_documents", default_return=[])
def load_documents(folder_path: str) -> List[Dict[str, Any]]:
    """Enhanced document loading with production features"""
    folder_path = Path(folder_path)
    
    if not folder_path.exists():
        logger.warning(f"Data folder does not exist: {folder_path}")
        return []
    
    logger.info(f"Loading documents from {folder_path}")
    start_time = time.time()
    
    processor = DocumentProcessor()
    documents = []
    processed_count = 0
    error_count = 0
    
    # Get all supported files
    supported_files = []
    for file_path in folder_path.iterdir():
        if file_path.is_file() and file_path.suffix.lower() in config.ALLOWED_EXTENSIONS:
            supported_files.append(file_path)
    
    if not supported_files:
        logger.info("No supported files found")
        return []
    
    logger.info(f"Found {len(supported_files)} supported files")
    
    for file_path in supported_files:
        try:
            file_result = processor.process_file(file_path)
            
            if file_result:
                text_content = file_result['content']
                
                # Create chunks for large documents
                if len(text_content) > config.CHUNK_SIZE * 2:
                    chunks = processor.chunk_text(text_content)
                    
                    for i, chunk in enumerate(chunks):
                        chunk_name = f"{file_path.name} (Part {i+1}/{len(chunks)})" if len(chunks) > 1 else file_path.name
                        documents.append({
                            'filename': chunk_name,
                            'content': chunk,
                            'original_filename': file_path.name,
                            'chunk_index': i,
                            'total_chunks': len(chunks),
                            'file_info': file_result['file_info']
                        })
                else:
                    # Keep smaller documents as single units
                    documents.append({
                        'filename': file_path.name,
                        'content': text_content,
                        'original_filename': file_path.name,
                        'chunk_index': 0,
                        'total_chunks': 1,
                        'file_info': file_result['file_info']
                    })
                
                processed_count += 1
        
        except Exception as e:
            error_count += 1
            logger.error(f"Failed to process {file_path.name}: {e}")
            continue
    
    total_time = time.time() - start_time
    total_content = sum(len(doc['content']) for doc in documents)
    
    logger.info(f"Document loading completed in {total_time:.2f}s", extra={
        'total_files': len(supported_files),
        'processed_files': processed_count,
        'error_count': error_count,
        'total_documents': len(documents),
        'total_content_chars': total_content
    })
    
    return documents